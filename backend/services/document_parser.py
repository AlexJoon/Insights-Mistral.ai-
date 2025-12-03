"""
Document parser service for extracting text from various file formats.
Uses Strategy Pattern to support multiple document types.
"""
from abc import ABC, abstractmethod
from typing import Optional, Dict
from dataclasses import dataclass
from pathlib import Path
import logging
import re

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    """
    Represents a parsed document with extracted content and metadata.

    Attributes:
        content: Extracted text content
        metadata: Document metadata (filename, course_code, etc.)
        num_pages: Number of pages (if applicable)
        file_type: Original file type (pdf, docx, txt)
    """
    content: str
    metadata: Dict
    num_pages: Optional[int] = None
    file_type: Optional[str] = None


class DocumentParseError(Exception):
    """Raised when document parsing fails."""
    pass


class DocumentParserStrategy(ABC):
    """
    Abstract base class for document parsing strategies.

    Design Pattern: Strategy Pattern
    Purpose: Different algorithms for different file types
    """

    @abstractmethod
    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file."""
        pass

    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        """Parse the document and extract text."""
        pass


class PDFParserStrategy(DocumentParserStrategy):
    """
    PDF document parser using PyPDF2.

    Handles:
    - Text extraction from PDFs
    - Multi-page documents
    - Basic metadata extraction
    """

    def can_parse(self, file_path: str) -> bool:
        """Check if file is a PDF."""
        return file_path.lower().endswith('.pdf')

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF file and extract text.

        Args:
            file_path: Path to PDF file

        Returns:
            ParsedDocument with extracted content

        Raises:
            DocumentParseError: If parsing fails
        """
        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                # Extract text from all pages
                text_content = []
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append(text)

                # Combine all pages
                full_text = "\n\n".join(text_content)

                # Clean up text (remove excessive whitespace)
                full_text = self._clean_text(full_text)

                # Extract basic metadata
                metadata = self._extract_metadata(full_text, file_path)
                metadata['num_pages'] = num_pages

                logger.info(f"Parsed PDF: {file_path} ({num_pages} pages, {len(full_text)} chars)")

                return ParsedDocument(
                    content=full_text,
                    metadata=metadata,
                    num_pages=num_pages,
                    file_type='pdf'
                )

        except ImportError:
            raise DocumentParseError(
                "PyPDF2 is not installed. Install with: pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            raise DocumentParseError(f"PDF parsing failed: {e}")

    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove page numbers (common pattern)
        text = re.sub(r'\n\s*\d+\s*\n', '\n', text)
        return text.strip()

    def _extract_metadata(self, text: str, file_path: str) -> Dict:
        """
        Extract metadata from syllabus content.

        Attempts to find:
        - Course code (e.g., CS101, MATH 201)
        - Course title
        - Instructor name
        - Semester/term
        """
        metadata = {
            'source_file': Path(file_path).name
        }

        # Try to extract course code (e.g., CS101, MATH 201, CSE-101)
        course_code_patterns = [
            r'\b([A-Z]{2,4}[-\s]?\d{3,4}[A-Z]?)\b',  # CS101, MATH 201, CSE-101
            r'\b([A-Z]{2,4}\d{3,4}[A-Z]?)\b',  # CS101, MATH201
        ]

        for pattern in course_code_patterns:
            match = re.search(pattern, text[:500])  # Search in first 500 chars
            if match:
                metadata['course_code'] = match.group(1).replace(' ', '').replace('-', '')
                break

        # Try to extract instructor (look for "Instructor:", "Professor:", etc.)
        instructor_patterns = [
            r'(?:Instructor|Professor|Teacher):\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            r'(?:Instructor|Professor|Teacher)\s*[-–]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
        ]

        for pattern in instructor_patterns:
            match = re.search(pattern, text[:1000], re.IGNORECASE)
            if match:
                metadata['instructor'] = match.group(1).strip()
                break

        # Try to extract semester (e.g., "Fall 2024", "Spring 2025")
        semester_pattern = r'\b(Fall|Spring|Summer|Winter)\s+\d{4}\b'
        match = re.search(semester_pattern, text[:1000], re.IGNORECASE)
        if match:
            metadata['semester'] = match.group(0)

        return metadata


class DOCXParserStrategy(DocumentParserStrategy):
    """
    DOCX document parser using python-docx.

    Handles:
    - Text extraction from Word documents
    - Paragraph preservation
    - Table content
    """

    def can_parse(self, file_path: str) -> bool:
        """Check if file is a DOCX."""
        return file_path.lower().endswith('.docx')

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX file and extract text.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument with extracted content

        Raises:
            DocumentParseError: If parsing fails
        """
        try:
            import docx

            doc = docx.Document(file_path)

            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))

            # Combine all content
            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n" + "\n".join(tables_text)

            # Extract metadata (reuse PDF metadata extractor)
            pdf_parser = PDFParserStrategy()
            metadata = pdf_parser._extract_metadata(full_text, file_path)

            logger.info(f"Parsed DOCX: {file_path} ({len(full_text)} chars)")

            return ParsedDocument(
                content=full_text,
                metadata=metadata,
                file_type='docx'
            )

        except ImportError:
            raise DocumentParseError(
                "python-docx is not installed. Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise DocumentParseError(f"DOCX parsing failed: {e}")


class TXTParserStrategy(DocumentParserStrategy):
    """
    Plain text parser for .txt and .md files.

    Simple text file reading with encoding detection.
    """

    def can_parse(self, file_path: str) -> bool:
        """Check if file is TXT or MD."""
        return file_path.lower().endswith(('.txt', '.md'))

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse text file.

        Args:
            file_path: Path to text file

        Returns:
            ParsedDocument with extracted content

        Raises:
            DocumentParseError: If parsing fails
        """
        try:
            # Try UTF-8 first, fall back to latin-1
            encodings = ['utf-8', 'latin-1', 'cp1252']

            content = None
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise DocumentParseError(f"Could not decode file with any encoding: {encodings}")

            # Extract metadata
            pdf_parser = PDFParserStrategy()
            metadata = pdf_parser._extract_metadata(content, file_path)

            logger.info(f"Parsed TXT: {file_path} ({len(content)} chars)")

            return ParsedDocument(
                content=content,
                metadata=metadata,
                file_type='txt'
            )

        except Exception as e:
            logger.error(f"Failed to parse TXT {file_path}: {e}")
            raise DocumentParseError(f"TXT parsing failed: {e}")


class DocumentParser:
    """
    Main document parser that delegates to appropriate strategy.

    Design Pattern: Strategy Pattern + Facade Pattern
    Purpose: Unified interface for parsing different document types

    Usage:
        parser = DocumentParser()
        doc = parser.parse("syllabus.pdf")
        print(doc.content)
        print(doc.metadata)
    """

    def __init__(self):
        """Initialize parser with all available strategies."""
        self.strategies = [
            PDFParserStrategy(),
            DOCXParserStrategy(),
            TXTParserStrategy()
        ]

        logger.info(f"DocumentParser initialized with {len(self.strategies)} strategies")

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document using the appropriate strategy.

        Args:
            file_path: Path to the document file

        Returns:
            ParsedDocument with extracted content

        Raises:
            DocumentParseError: If no parser can handle the file or parsing fails
        """
        # Find appropriate strategy
        for strategy in self.strategies:
            if strategy.can_parse(file_path):
                logger.info(f"Using {strategy.__class__.__name__} for {file_path}")
                return strategy.parse(file_path)

        # No strategy found
        supported_types = ['.pdf', '.docx', '.txt', '.md']
        raise DocumentParseError(
            f"Unsupported file type: {file_path}. "
            f"Supported types: {', '.join(supported_types)}"
        )

    def can_parse(self, file_path: str) -> bool:
        """
        Check if any strategy can parse this file.

        Args:
            file_path: Path to the document file

        Returns:
            True if file can be parsed, False otherwise
        """
        return any(strategy.can_parse(file_path) for strategy in self.strategies)
